from backend.services.parsing.docx_parser import DocxParser


def test_heading_levels(tmp_path):
    from docx import Document as DocxDocument

    d = DocxDocument()
    d.add_heading("Chapter 1", level=1)
    d.add_paragraph("Some body text.")
    d.add_heading("Section 1.1", level=2)
    d.add_paragraph("More text.")
    path = tmp_path / "test.docx"
    d.save(str(path))
    parser = DocxParser()
    elements = parser.parse(str(path))
    headings = [e for e in elements if e.element_type == "heading"]
    assert headings[0].content == "Chapter 1"
    assert headings[0].heading_level == 1
    assert headings[1].content == "Section 1.1"
    assert headings[1].heading_level == 2


def test_table_extraction(tmp_path):
    from docx import Document as DocxDocument

    d = DocxDocument()
    d.add_heading("Table Section", level=1)
    table = d.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "A"
    table.cell(0, 1).text = "B"
    table.cell(1, 0).text = "1"
    table.cell(1, 1).text = "2"
    path = tmp_path / "table.docx"
    d.save(str(path))
    parser = DocxParser()
    elements = parser.parse(str(path))
    tables = [e for e in elements if e.element_type == "table"]
    assert len(tables) == 1
    assert "A" in tables[0].content
