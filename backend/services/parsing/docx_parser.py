"""
DOCX 解析器模块

使用 python-docx 库解析 Word 文档，提取：
- 标题段落（Heading 1-6）
- 普通段落
- 表格（转换为 Markdown 格式）

返回结构化元素列表。
"""

from datetime import datetime, timezone

from docx import Document as DocxDocument
from backend.services.parsing.base import BaseParser, StructuredElement


class DocxParser(BaseParser):
    """
    Word DOCX 文档解析器

    解析 DOCX 文件中的段落和表格，支持识别标题级别。
    """

    def parse(self, filepath: str) -> list[StructuredElement]:
        """
        解析 DOCX 文件

        Args:
            filepath: DOCX 文件路径

        Returns:
            结构化元素列表
        """
        doc = DocxDocument(filepath)
        created = doc.core_properties.created
        if created and isinstance(created, datetime):
            created_at = created.astimezone(timezone.utc).isoformat()
        else:
            created_at = None
        elem_meta = {"created_at": created_at} if created_at else {}

        result: list[StructuredElement] = []
        for p in doc.paragraphs:
            style = p.style.name
            text = p.text.strip()
            if not text:
                continue
            if style.startswith("Heading "):
                level = int(style.split()[1])
                result.append(
                    StructuredElement(
                        content=text,
                        element_type="heading",
                        heading_level=min(level, 6),
                        metadata=elem_meta,
                    )
                )
            else:
                result.append(
                    StructuredElement(content=text, element_type="paragraph",
                        metadata=elem_meta)
                )
        for tbl in doc.tables:
            rows = []
            for row in tbl.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(" | ".join(cells))
            md_table = "\n".join(rows)
            result.append(StructuredElement(content=md_table, element_type="table",
                metadata=elem_meta))
        return result
